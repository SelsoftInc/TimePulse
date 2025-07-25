from fastapi import FastAPI, Header, HTTPException
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from azure.core.exceptions import HttpResponseError
from langchain_community.chat_models import AzureChatOpenAI
from langchain.prompts import PromptTemplate
from langchain_community.callbacks import get_openai_callback
from dotenv import load_dotenv
import os
import json
import tempfile
from docx import Document
import requests
from fastapi.responses import FileResponse
from fpdf import FPDF

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI app
app = FastAPI()

# Retrieve Azure OpenAI credentials from environment variables
azure_openai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
azure_openai_api_key = os.getenv("AZURE_OPENAI_API_KEY")
azure_openai_deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
azure_openai_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2023-06-01-preview")

if not azure_openai_endpoint or not azure_openai_api_key or not azure_openai_deployment_name:
    raise ValueError("Azure OpenAI endpoint, API key, and deployment name must be set in the .env file.")

# Initialize Azure OpenAI GPT model with LangChain
chat_model = AzureChatOpenAI(
    azure_deployment=azure_openai_deployment_name,
    api_version=azure_openai_api_version,
    temperature=0.7,
    max_tokens=None,
    max_retries=2
)

# Retrieve Azure Document Intelligence credentials
endpoint = os.getenv("DOCUMENTINTELLIGENCE_ENDPOINT")
key = os.getenv("DOCUMENTINTELLIGENCE_API_KEY")

if not endpoint or not key:
    raise ValueError("Azure endpoint and key must be set in the .env file.")

document_intelligence_client = DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))

BASE_DIR = "/home/nibi/Desktop/Projects/Selsoft/"
DOCX_PATH = os.path.join(BASE_DIR, "temp.docx")
PDF_PATH = os.path.join(BASE_DIR, "output.pdf")

# Define the directory for saving files
BASE_DIR = "/home/nibi/Desktop/Projects/Selsoft/"
DOCX_PATH = os.path.join(BASE_DIR, "temp.docx")
PDF_PATH = os.path.join(BASE_DIR, "output.pdf")

@app.post("/convert-docx-to-pdf")
def convert_docx_to_pdf(docx_url: str = Header(...)):
    try:
        # Ensure the directory exists
        os.makedirs(BASE_DIR, exist_ok=True)

        # Fetch the DOCX file from the URL
        response = requests.get(docx_url)
        if response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to fetch the DOCX file.")

        # Save the DOCX file locally
        with open(DOCX_PATH, "wb") as docx_file:
            docx_file.write(response.content)

        # Read the DOCX file
        document = Document(DOCX_PATH)
        content = "\n".join(paragraph.text for paragraph in document.paragraphs)

        # Convert DOCX content to PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for line in content.splitlines():
            pdf.cell(0, 10, txt=line, ln=True)

        # Save the PDF and verify
        pdf.output(PDF_PATH)
        if not os.path.exists(PDF_PATH):
            raise RuntimeError(f"Failed to create PDF at path {PDF_PATH}.")

        # Return the PDF file as a response
        return FileResponse(PDF_PATH, media_type="application/pdf", filename="output.pdf")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
   
@app.get("/analyze-document")
async def analyze_document(document_url: str = Header(...)):
    """
    Analyze a document from the provided URL using Azure Document Intelligence service.
    The URL should be passed as a header: `document-url`
    If the document is a DOCX file, it will be converted to PDF before analysis.
    """
    try:
        if document_url.endswith(".docx"):
            local_docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode="wb+").name
            local_pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name

            # Download DOCX file
            response = requests.get(document_url)
            if response.status_code == 200:
                with open(local_docx_path, "wb") as docx_file:
                    docx_file.write(response.content)
            else:
                raise HTTPException(status_code=400, detail="Failed to download the DOCX file.")

            # Convert DOCX to PDF
            document = Document(local_docx_path)
            with open(local_pdf_path, "w") as pdf_file:
                for paragraph in document.paragraphs:
                    pdf_file.write(paragraph.text + "\n")

            # Analyze PDF content
            with open(local_pdf_path, "rb") as f:
                poller = document_intelligence_client.begin_analyze_document(
                    model_id="prebuilt-read",
                    document=f.read(),
                    content_type="application/octet-stream"
                )
                result = poller.result()



        else:
            poller = document_intelligence_client.begin_analyze_document(
            "prebuilt-read",
            {"urlSource": document_url}
            )
            result = poller.result()
       
        # Extract content from the analyzed result
        content = "\n".join(paragraph.content for paragraph in result.paragraphs)

        prompt_template = PromptTemplate(
            input_variables=["document"],
            template="""
            You are a language model tasked with analyzing and summarizing the following time sheet document content extracted using OCR:
            {document}
            Provide a JSON response consisting of the name of the vendor, their dates, and hours in the below format.
            {{
                "Entry": [
                    {{
                        "Vendor_Name": Name,
                        "Total Hours": Hours,
                        "Duration": "From date to end date"
                    }}
                ]
            }}
            """
        )

        with get_openai_callback() as cb:
            # Generate a response
            response = chat_model.invoke(
                messages=[{"role": "user", "content": prompt_template.format(document=content)}]
            )

            cost = cb.total_cost
            feedback = response.content
            json_start = feedback.find('{')
            json_end = feedback.rfind('}') + 1

            json_string = feedback[json_start:json_end]

            results = json.loads(json_string)

        return {"status": "Processed successfully", "results": results, "cost": f"Total Cost (USD): ${format(cost, '.6f')}"}

    except HttpResponseError as error:
        if error.error:
            return {"error": error.error.code, "message": error.error.message}
        return {"error": "HttpResponseError", "message": str(error)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
